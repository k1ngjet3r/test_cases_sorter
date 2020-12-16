from openpyxl import load_workbook
from docx import Document

sw_build = 'Mainline,  W94I-W157.2.1-QIH22B-220 Dev-signed'


def ticket_gen(input_file, sw):
    # Initialized the excel file which containing the detail
    sheet = load_workbook(str(input_file)).active

    for row in sheet.iter_rows(max_col=11, values_only=True):
        row_data = [str(value) for value in row]
        doc_name = row_data[-1]
        if row_data[0] == 'Summary':
            pass
        elif row_data[0] == 'None':
            break
        else:
            # Open up a blank document with the default template
            ticket = Document()

            # Summary section
            summary = ticket.add_paragraph('')
            summary_title = summary.add_run('[Summary] ')
            summary_title.bold = True
            summary.add_run(row_data[0])

            # Precondition
            precondiction = ticket.add_paragraph('')
            precondiction.add_run('[Precondition]').bold = True
            # Testing Type
            testing_type = ticket.add_paragraph('')
            testing_type.add_run('Testing Type: ').bold = True
            testing_type.add_run(row_data[1])
            # Connected Devices
            connected_devices = ticket.add_paragraph('')
            connected_devices.add_run('Connected Devices ').bold = True
            connected_devices.add_run(row_data[2])
            ticket.add_paragraph(row_data[3])
            # Test Steps
            test_steps = ticket.add_paragraph('')
            test_steps.add_run('[Test Steps] ').bold = True
            ticket.add_paragraph(row_data[4])

            # Expected Result
            expected_result = ticket.add_paragraph('')
            expected_result.add_run('[Expected Result] ').bold = True
            ticket.add_paragraph(row_data[5])

            # Actual Result
            actual = ticket.add_paragraph('')
            actual.add_run('[Actual Result] ').bold = True
            ticket.add_paragraph(row_data[6])

            # Reproduced Rate
            repo = ticket.add_paragraph('')
            repo.add_run('[Reproduced Rate] ').bold = True
            ticket.add_paragraph('10/10')

            # Spec
            spec = ticket.add_paragraph('')
            spec.add_run('[Spec Reference]').bold = True
            ticket.add_paragraph(row_data[7])

            # time
            time = ticket.add_paragraph('')
            time.add_run('[Occurrence Time] ').bold = True
            time.add_run(row_data[8])

            # Comments
            comments = ticket.add_paragraph('')
            comments.add_run('[Comments]').bold = True
            ticket.add_paragraph(row_data[9])

            # Hardware
            hardware = ticket.add_paragraph('')
            hardware.add_run('[Hardware Info]').bold = True
            csm = ticket.add_paragraph('')
            csm.add_run('CSM: ').bold = True
            csm.add_run('GB MY22 CSM3.7 PV01H, High, NA')

            # Related cases
            rel_case = ticket.add_paragraph('')
            rel_case.add_run('[Found by running test case]').bold = True
            ticket.add_paragraph(row_data[10])

            # Other
            sw = ticket.add_paragraph('')
            sw.add_run('SW build info: ').bold = True
            sw.add_run(sw_build)

            contact = ticket.add_paragraph('')
            contact.add_run('Contact Phone Number: ').bold = True
            contact.add_run('+886 966603203')

            sub = ticket.add_paragraph('')
            sub.add_run('Submitter: ').bold = True
            sub.add_run('<<<rick.weng@cienet.com>>')

            ticket.save('{}.docx'.format(doc_name))


ticket_gen('Defect_Tickets1.xlsx', sw_build)
